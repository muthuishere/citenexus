"""Captions reach ask() as their own cited passage — end to end.

Closes docs/CONTENT-COVERAGE-2026-07-08.md row 6's CAPTION half (footnotes
are a separate, still-open question — see the coverage doc for why: python-docx
has no footnote-part API, and markdown-it-py's footnote syntax needs an
extra plugin dependency not currently in the project — out of scope here,
not silently claimed).

extract/html.py now captures <figcaption> (a figure's real caption — distinct
from a vision model's own generated short_caption) and <table><caption>.
extract/docx.py needed NO change: a paragraph styled "Caption" (Word's
built-in caption style) was already just an unstyled BlockKind.paragraph to
_heading_level — verified live, not just asserted, same pattern as row 3.
"""

from __future__ import annotations

from pathlib import Path

from docx import Document

from citenexus import CiteNexus
from citenexus.answer.result import Decision
from citenexus.testing import FakeEmbedding, FakeLLM

_HTML = """
<html><body>
<p>Unrelated narrative text.</p>
<figure><img src="chart.png">
<figcaption>Figure 1: revenue rose 20 percent in Q4.</figcaption>
</figure>
<table><caption>Table 1: headcount was 42 employees at year end.</caption>
<tr><th>Team</th><th>Count</th></tr><tr><td>Eng</td><td>12</td></tr>
</table>
</body></html>
"""


def test_html_figcaption_is_retrieved_and_cited(tmp_path: Path) -> None:
    html_path = tmp_path / "report.html"
    html_path.write_text(_HTML)
    rag = CiteNexus(tmp_path / "store", embedder=FakeEmbedding(), generator=FakeLLM())
    result = rag.ingest(html_path, document_id="report-html")
    assert result.status == "ingested"

    answer = rag.ask("By what percent did revenue rise in Q4?")

    assert answer.evidence.decision is Decision.answered
    assert answer.claims[0].supported
    assert "20 percent" in answer.sources[0].passage


def test_html_table_caption_is_retrieved_and_cited(tmp_path: Path) -> None:
    html_path = tmp_path / "report.html"
    html_path.write_text(_HTML)
    rag = CiteNexus(tmp_path / "store", embedder=FakeEmbedding(), generator=FakeLLM())
    result = rag.ingest(html_path, document_id="report-html")
    assert result.status == "ingested"

    answer = rag.ask("What was the headcount at year end?")

    assert answer.evidence.decision is Decision.answered
    assert answer.claims[0].supported
    assert "42 employees" in answer.sources[0].passage


def test_docx_caption_style_is_retrieved_and_cited(tmp_path: Path) -> None:
    doc = Document()
    doc.add_paragraph("Unrelated narrative text.")
    caption = doc.add_paragraph("Figure 1: revenue rose 20 percent in Q4.")
    caption.style = doc.styles["Caption"]
    docx_path = tmp_path / "report.docx"
    doc.save(str(docx_path))

    rag = CiteNexus(tmp_path / "store", embedder=FakeEmbedding(), generator=FakeLLM())
    result = rag.ingest(docx_path, document_id="report-docx")
    assert result.status == "ingested"

    answer = rag.ask("By what percent did revenue rise in Q4?")

    assert answer.evidence.decision is Decision.answered
    assert answer.claims[0].supported
    assert "20 percent" in answer.sources[0].passage
