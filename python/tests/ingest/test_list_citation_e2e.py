"""List items reach ask() as the cited passage — end to end.

Closes docs/CONTENT-COVERAGE-2026-07-08.md row 3 for HTML (the genuinely
broken one — <ul>/<ol>/<li> was entirely outside HtmlExtractor's selector,
content dropped, not even flattened). Markdown and DOCX were re-verified
against the live extractors and, contrary to the doc's prior note, already
reach the evidence store today (Markdown: markdown-it's "hidden" paragraph_open
tokens inside tight list items still hit the existing paragraph_open branch;
DOCX: a "List Bullet"-styled paragraph is just an unstyled BlockKind.paragraph
to _heading_level) — both proven here too, not just asserted.
"""

from __future__ import annotations

from pathlib import Path

from docx import Document

from citenexus import CiteNexus
from citenexus.answer.result import Decision
from citenexus.testing import FakeEmbedding, FakeLLM


def test_html_list_item_is_retrieved_and_cited(tmp_path: Path) -> None:
    html_path = tmp_path / "policy.html"
    html_path.write_text(
        "<html><body><p>Unrelated narrative text.</p>"
        "<ul><li>Annual leave: 24 days per year.</li>"
        "<li>Sick leave: 12 days per year.</li></ul></body></html>"
    )
    rag = CiteNexus(tmp_path / "store", embedder=FakeEmbedding(), generator=FakeLLM())
    result = rag.ingest(html_path, document_id="leave-policy-html")
    assert result.status == "ingested"

    answer = rag.ask("How many days of annual leave per year?")

    assert answer.evidence.decision is Decision.answered
    assert answer.claims[0].supported
    assert "24 days" in answer.sources[0].passage


def test_markdown_list_item_is_retrieved_and_cited(tmp_path: Path) -> None:
    md_path = tmp_path / "policy.md"
    md_path.write_text(
        "Unrelated narrative text.\n\n"
        "- Annual leave: 24 days per year.\n"
        "- Sick leave: 12 days per year.\n"
    )
    rag = CiteNexus(tmp_path / "store", embedder=FakeEmbedding(), generator=FakeLLM())
    result = rag.ingest(md_path, document_id="leave-policy-md")
    assert result.status == "ingested"

    answer = rag.ask("How many days of annual leave per year?")

    assert answer.evidence.decision is Decision.answered
    assert answer.claims[0].supported
    assert "24 days" in answer.sources[0].passage


def test_docx_list_item_is_retrieved_and_cited(tmp_path: Path) -> None:
    doc = Document()
    doc.add_paragraph("Unrelated narrative text.")
    doc.add_paragraph("Annual leave: 24 days per year.", style="List Bullet")
    doc.add_paragraph("Sick leave: 12 days per year.", style="List Bullet")
    docx_path = tmp_path / "policy.docx"
    doc.save(str(docx_path))

    rag = CiteNexus(tmp_path / "store", embedder=FakeEmbedding(), generator=FakeLLM())
    result = rag.ingest(docx_path, document_id="leave-policy-docx")
    assert result.status == "ingested"

    answer = rag.ask("How many days of annual leave per year?")

    assert answer.evidence.decision is Decision.answered
    assert answer.claims[0].supported
    assert "24 days" in answer.sources[0].passage
